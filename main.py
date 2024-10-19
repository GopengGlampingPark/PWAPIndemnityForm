from flask import Flask, url_for, render_template, redirect, request, session, jsonify, copy_current_request_context, send_file
from flask_session import Session
from flask_mail import Mail, Message
import base64
from docx import Document
from docx.shared import Inches
from io import BytesIO
from PIL import Image
import gspread
from google.oauth2.service_account import Credentials
from oauth2client.service_account import ServiceAccountCredentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import logging
from threading import Thread, Lock
import json
import os
from dotenv import load_dotenv
import sqlite3
import redis

app = Flask(__name__)

# Configure Flask-Session
app.config['SESSION_TYPE'] = 'filesystem'  # This can be changed based on your needs
app.config['SESSION_PERMANENT'] = False
app.config['SESSION_USE_SIGNER'] = True
app.secret_key = "GGPINDEMNITYFORMFROMFUAD"
Session(app)

# add comments pleaseeeeee :D
load_dotenv('.env')  # Adjust path if necessary

# Configure Flask-Mail with fallback values
app.config['MAIL_SERVER'] = os.getenv('MAIL_SERVER', 'smtp.gmail.com')  # Default to Gmail SMTP
app.config['MAIL_PORT'] = int(os.getenv('MAIL_PORT', 587))  # Default to port 587
app.config['MAIL_USE_TLS'] = os.getenv('MAIL_USE_TLS', 'true') == 'true'
app.config['MAIL_USE_SSL'] = os.getenv('MAIL_USE_SSL', 'false') == 'true'
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_DEFAULT_SENDER')

# Initialize Flask-Mail
mail = Mail(app)

# redis for counter
redis_host = os.getenv('REDIS_HOST')
redis_port = os.getenv('REDIS_PORT')
redis_password = os.getenv('REDIS_PASSWORD')

r = redis.Redis(
	host=redis_host,
	port=redis_port,
	password=redis_password,
	ssl=True
)

logging.basicConfig(
	level=logging.DEBUG,  # Set the logging level
	format='%(asctime)s - %(levelname)s - %(message)s',  # Log format
)

# Google Sheets credentials and initialization
SCOPES = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive.file']

#### RENDER #########################################################################################################################################
SERVICE_ACCOUNT_INFO = json.loads(os.environ.get('GOOGLE_SERVICE_ACCOUNT_JSON'))
credentials = ServiceAccountCredentials.from_json_keyfile_dict(SERVICE_ACCOUNT_INFO, scopes=SCOPES)

#### LOCAL #########################################################################################################################################
### SERVICE_ACCOUNT_INFO = 'ggp-indemnity-form-5ae555d2a987.json'
### credentials = ServiceAccountCredentials.from_json_keyfile_name(SERVICE_ACCOUNT_INFO, scopes=SCOPES)

client = gspread.authorize(credentials)

# Open the spreadsheet
SPREADSHEET_ID = '1UGBl85ALPONDjI4D4JXkVWmopQnPruTjXzGJxCZIqZQ'
spreadsheet = client.open_by_key(SPREADSHEET_ID)

# Initialize a lock for thread-safe ID generation
lock = Lock()

######################################################################################################################## WEB APPPP

@app.route("/genbookid", methods=["POST", "GET"])
def genbookid():
	if request.method == "POST":
		C_bookingID = generate_sequential_id("GroupInformation", "PWAPBOOKID_")
		session['bookingName'] = request.form.get('bookingName')  # Ensure 'groupName' comes from the form
		session['C_bookingID'] = C_bookingID
		bookingName = session.get("bookingName")
		print(f"booking name : {bookingName}")

		@copy_current_request_context
		def rungroupintosheet():
			try:
				print(f"running group threads")
				insert_group_to_sheet()  # Make sure this function exists and handles its logic
			except Exception as e:
				# Log exceptions using Flask's logging
				logging.error(f"Error in background thread: {str(e)}")

		thread = Thread(target=rungroupintosheet)
		thread.start()
		return render_template("bookingID.html", C_bookingID=C_bookingID, bookingName=bookingName)
	return render_template("bookingID.html")

def insert_group_to_sheet():
	bookingName = session.get("bookingName")
	C_bookingID = session.get("C_bookingID")

	if C_bookingID:
		group_info = [[
			C_bookingID,
			bookingName
		]]
		print("Group INFO : ", group_info)
		append_to_sheet('GroupInformation', group_info)

@app.route("/", methods=["POST", "GET"])
def indemnityform():
	if request.method == "POST":
		formfield = ["fullname", "gender", "age", "NRIC", "email", "contact", "lineaddress", "postcode", "state", "country", "nationality", "race", "staynot", "actamount", "bookingID"]
		for field in formfield:
			session[field] = request.form[field]

		if session.get("age") == "under12" or session.get("age") == "13-17":
			return redirect(url_for("under18"))
		else:
			if session["staynot"] == "stayinguest":
				return redirect(url_for("checkinout"))
			else:
				return redirect(url_for("healthdeclare"))
	else:
		return render_template("form.html")

@app.route("/under18", methods=["POST", "GET"])
def under18():
	if request.method == 'POST':
		try:
			gsignature_data = request.form.get('gsignature')
			session['gsignature'] = gsignature_data
			under18_fields = ["acknowledgement", "gname", "gcontact"]
			for field in under18_fields:
				session[field] = request.form[field]

			print("Received signature data:", session.get("gsignature"))

			if session.get("staynot") == "stayinguest":
				next_url = url_for("checkinout")
			else:
				next_url = url_for("healthdeclare")

			return jsonify({'next_url': next_url}), 200  # Only include next_url

		except Exception as e:
			print("Error:", e)
			return jsonify({'error': 'Failed to receive signature'}), 400  # Minimal error response

	elif "age" in session:
		if session.get("age") == "under12" or session.get("age") == "13-17":
			return render_template("under18.html")
	
	else:
		return redirect(url_for("indemnityform"))

@app.route("/checkinout", methods=["POST", "GET"])
def checkinout():
	if request.method == 'POST':
		checkinout_fields = ["checkin", "checkout"]
		for field in checkinout_fields:
			session[field] = request.form[field]
		return redirect(url_for("healthdeclare"))

	elif session["staynot"] == "stayinguest":
		return render_template("checkinout.html")
	elif session["staynot"] == "nonstayinguest":
		return redirect(url_for("healthdeclare"))
	else:
		return redirect(url_for("indemnityform"))

@app.route("/healthdeclare", methods=["POST", "GET"])
def healthdeclare():
	if request.method == "POST":
		health_fields = request.form.getlist("health_fields[]")
		session["health_fields"] = health_fields
		return redirect(url_for("activity"))
	elif "fullname" in session:
		return render_template("healthdeclare.html")
	else:
		return redirect(url_for("indemnityform"))

@app.route("/activity", methods=["POST", "GET"])
def activity():
	# Retrieve actamount from the session; it should be set during a previous interaction
	actamount = int(session.get("actamount"))  # No default value

	if request.method == "POST":
		# Initialize a list to hold activities
		activities = []

		for i in range(actamount):
			activity_name = request.form.get(f'activity{i}')  # e.g., 'WWR'
			activity_date = request.form.get(f'act_date{i}')  # e.g., '2424-04-24'
			activity_time = request.form.get(f'act_time{i}')  # e.g., 'PM230'

			# Only add activity if all required fields are filled
			if activity_name and activity_date and activity_time:
				activities.append({
					'AN': activity_name,  # Activity Name
					'AD': activity_date,	# Activity Date
					'AT': activity_time	 # Activity Time
				})
			else:
				print(f"Missing data for activity {i}: Name: {activity_name}, Date: {activity_date}, Time: {activity_time}")

		# Store the list of activities and actamount in the session
		session['activities'] = activities  
		session['actamount'] = actamount  # Store actamount in the session
		print("Stored activities:", session.get('activities'))

		return redirect(url_for("tnc"))  # Redirect to the Terms and Conditions page

	elif "fullname" in session:
		# Ensure actamount is available when rendering
		if actamount is None:
			return redirect(url_for("indemnityform"))  # Redirect if actamount is not set

		return render_template("activity.html", actamount=actamount)  # Render the activity form page
	else:
		return redirect(url_for("indemnityform"))

@app.route("/tnc", methods=["POST", "GET"])
def tnc():
	if request.method == "POST":
		try:
			psignature_data = request.form.get('psignature')
			session["psignature"] = psignature_data
			print("psignature received:", session.get("psignature"))
			next_url = url_for("thankyou")
			return jsonify({'next_url': next_url}), 200  # Only include next_url

		except Exception as e:
			print(f"Error processing psignature: {str(e)}")
			return jsonify({"status": "error"}), 400  # Minimal error response

	elif "fullname" in session:
		return render_template("tnc.html")
	else:
		return redirect(url_for("indemnityform"))

@app.route("/thankyou", methods=["POST", "GET"])
def thankyou():
	@copy_current_request_context
	def runinserttosheet():
		try:
			print(f"running threads")
			insert_to_sheet()  # Make sure this function exists and handles its logic
		except Exception as e:
			# Log exceptions using Flask's logging
			logging.error(f"Error in background thread: {str(e)}")

	thread = Thread(target=runinserttosheet)
	thread.start()

	return render_template("thankyou.html")

######################################################################################################################## SHEET RELATED
def insert_to_sheet():
	C_UniqueID = generate_sequential_id("Client_Information" , "PWAPID_")
	
	C_bookingID = session["bookingID"]  # Use the existing bookingID from the session if it's not None or empty

	print(f"sending E-Cert")
	
	gsignature_data = session.get('gsignature')
	psignature_data = session.get('psignature')
	print(psignature_data)
	# Check and save gsignature
	if gsignature_data:
		gsigsaved = save_signature(session.get("gsignature"))  # Save the gsignature
		print("gsigsaved:", gsigsaved)  # Show the saved value
	else:
		print("gsigsaved not saved")

	# Check and save psignature
	if psignature_data:
		psigsaved = save_signature(session.get("psignature"))  # Save the psignature
		print("psigsaved:", psigsaved)  # Show the saved value
	else:
		print("psigsaved not saved")

	# Prepare and append client info
	clientinfo = [[
		C_UniqueID,
		C_bookingID,
		session.get("fullname"),
		session.get("gender"),
		session.get("age"),
		session.get("NRIC"),
		session.get("email"),
		session.get("contact"),
		session.get("lineaddress"),
		session.get("postcode"),
		session.get("state"),
		session.get("country"),
		session.get("nationality"),
		session.get("race"),
		session.get("staynot")
	]]
	print("cleintinfo : ", clientinfo)
	append_to_sheet('Client_Information', clientinfo)

	if session.get("age") == "under12" or session.get("age") == "13-17":
		under18info = [[
			C_UniqueID,
			session.get("acknowledgement"),
			session.get("gname"),
			session.get("gcontact"),
			gsigsaved
		]]
		print("under18info : ", under18info)
		append_to_sheet('Under18', under18info)

	if session.get("staynot") == "stayinguest":
		bookingInfo = [[
			C_UniqueID,
			session.get("checkin"),
			session.get("checkout")
		]]
		print("staynotinfo : ", bookingInfo)
		append_to_sheet('BookingInformation', bookingInfo)

	if session.get("health_fields"):
		healthInfo = [[
			C_UniqueID,
			", ".join(session.get("health_fields", []))
		]]
		print("healthinfo : ", healthInfo)
		append_to_sheet('HealthInformation', healthInfo)

	# Append activities if available
	if session.get("activities"):
		activities_info = [
			[C_UniqueID, C_bookingID, activity['AN'], activity['AD'], activity['AT']]
			for activity in session['activities'] if isinstance(activity, dict)
		]

		if activities_info:
			print("Activity info : ", activities_info)
			append_to_sheet('ClientActivity', activities_info)
		else:
			print("No valid activities to append.")

	print("activity saved to sheet")

	if gsignature_data:
		terms_and_conditions_info = [[
			C_UniqueID,
			psigsaved,
			gsigsaved
		]]
		print("TNC INFO : ", terms_and_conditions_info)
		append_to_sheet('Signatures', terms_and_conditions_info)
	else:
		terms_and_conditions_info = [[
			C_UniqueID,
			psigsaved
		]]
		print("TNC INFO : ", terms_and_conditions_info)
		append_to_sheet('Signatures', terms_and_conditions_info)

	if session.get("health_fields"):
		health_emails(C_UniqueID)

	submit_form(C_UniqueID)

def append_to_sheet(sheet_name, values):
	# Open the specified sheet
	sheet = spreadsheet.worksheet(sheet_name)
	sheet.append_rows(values, value_input_option='RAW')
	logging.info(f"Appended values to {sheet_name}: {values}")

def save_signature(signature):
	if not signature:
		return None
	
	# Decode the base64 signature
	image_data = signature.split(",")[1]
	image = Image.open(BytesIO(base64.b64decode(image_data)))

	# Save the image to a BytesIO object
	image_bytes = BytesIO()
	image.save(image_bytes, format='PNG')
	image_bytes.seek(0)  # Move to the beginning of the BytesIO object
	
	# Create a file in Google Drive
	folder_id = '1dd6swFYEryXnsUn2khwJslrhCZW5qUHl'  # Specify your Google Drive folder ID here
	image_name = f"signature_{session.get('fullname', 'unknown')}.png"

	# Create a MediaIoBaseUpload object with the image data
	media = MediaIoBaseUpload(image_bytes, mimetype='image/png')

	# Create a file in the specified Google Drive folder
	drive_service = build('drive', 'v3', credentials=credentials)
	file_metadata = {
		'name': image_name,
		'parents': [folder_id]
	}
	
	try:
		file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
		driveimg = f"https://drive.google.com/file/d/{file.get('id')}"
		return driveimg
	except Exception as e:
		return {"status": "error", "message": str(e)}

######################################################################################################################## ID RELATED

# Step 2: Load the current ID from Redis
def load_id_storage(sheet_name):
	""" Retrieve the current ID for a given sheet name from Redis """
	current_id = r.get(sheet_name)  # Fetch current ID from Redis
	return int(current_id) if current_id else 0  # Return 0 if no entry exists

# Step 3: Save the updated ID back into Redis
def save_id_storage(sheet_name, new_id):
	""" Save the updated ID for the given sheet name in Redis """
	r.set(sheet_name, new_id)  # Store the new ID in Redis

# Step 4: Generate a new sequential ID
def generate_sequential_id(sheet_name, prefix):
	""" Generate a new sequential ID with the given prefix """
	# Load the current ID from Redis
	current_id = load_id_storage(sheet_name) + 1  # Increment the ID
	new_id = f"{prefix}{current_id}"
	# Save the updated ID back to Redis
	save_id_storage(sheet_name, current_id)
	return new_id

######################################################################################################################## HEALTH CONDITION EMAILS
def health_emails(clientID):

	print(f"sending health emails")
	activities_info = [
		[activity['AN'], activity['AD'], activity['AT']]
		for activity in session.get('activities', [])
	]

	print("setting up messages")
	clientName = session.get("fullname")
	CContact = session.get("contact")
	msg = Message(
			subject="Health Conditions Notification",
			sender=app.config['MAIL_USERNAME'],
			recipients=['pwadventurepark@gmail.com'],
		)

	print("setting up email body")
	email_body = (
		"Hello,\n\n"
		f"Please be informed that {clientName} with the id {clientID} has submitted their indemnity form for the upcoming adventure activity.\n"
		"They have declared certain medical conditions that may require your attention during the trip.\n\n"
		"Kindly review the details provided in the form and take any necessary precautions to ensure their safety and well-being throughout the activity. \n"
		"If you need additional information or support, feel free to reach out to the management team.\n\n"
		"Here are the details of the participants:\n\n"
		f"Name : {clientName}\n\n"
		f"Unique ID : {clientID}\n\n"
		f"Phone Number : {CContact}\n\n"
	)

	email_body += f"has these health conditions\n\n"
	email_body += f"\n ".join(session.get("health_fields", []))

	for activity in activities_info:
		email_body += f"Activity Name: {activity[0]}\n"
		email_body += f"Activity Date: {activity[1]}\n"
		email_body += f"Activity Time: {activity[2]}\n\n"

	email_body += f"Please refer to the google sheets for additional information\n\n"
	email_body += f"Thank you for your attention and cooperation\n\n"
	email_body += f"Best regards\n"
	email_body += f"Glamping Park Travel And Tour Sdn Bhd\n"
	msg.body = email_body

	try:
		print("Sending email")
		# Send the email
		mail.send(msg)
		print('email sent successfully!')
	except Exception as e:
		print(f"Error while sending email: {str(e)}")

########################################################################################################################
def insert_image_at_paragraph(paragraph, image_path):
	"""Insert an image at the location of the specified paragraph."""
	run = paragraph.add_run()
	run.add_picture(image_path, width=Inches(2))  # Adjust width if needed

def fix_base64_padding(base64_string):
	"""Fix base64 string padding if necessary."""
	missing_padding = len(base64_string) % 4
	if missing_padding:
		base64_string += '=' * (4 - missing_padding)  # Add the correct number of '='
	return base64_string

def is_valid_base64(base64_string):
	"""Check if the base64 string is valid."""
	try:
		base64.b64decode(base64_string, validate=True)  # Validate to avoid exceptions on padding
		return True
	except Exception:
		return False

def edit_docx_in_memory(file_path, replacements):
	"""Edit a .docx file by replacing placeholders with actual values and inserting a signature without saving to disk."""
	doc = Document(file_path)

	# Convert base64 signature to an image file if the psignature exists
	signature_base64 = session.get("psignature", "")
	print(f"Base64 Signature (Length: {len(signature_base64)})")  # Log the length of the signature

	signature_image_path = None

	if signature_base64:
		# Check if the base64 string contains the data URL prefix and remove it
		if signature_base64.startswith("data:image/png;base64,"):
			signature_base64 = signature_base64.split(",")[1]  # Remove the prefix

		# Fix any padding issues
		signature_base64 = fix_base64_padding(signature_base64)

		if is_valid_base64(signature_base64):
			try:
				# Decode the base64 string and save it as a temporary image file
				signature_data = base64.b64decode(signature_base64)
				signature_image_path = "temp_signature.png"

				# Save the decoded image data to a file (temporarily)
				with open(signature_image_path, "wb") as f:
					f.write(signature_data)
				print("Saved temp_signature")

				# Insert the signature image into the document
				for paragraph in doc.paragraphs:
					if "<<SIGNATURE>>" in paragraph.text:
						paragraph.text = paragraph.text.replace("<<SIGNATURE>>", "")  # Remove the placeholder
						insert_image_at_paragraph(paragraph, signature_image_path)

			except Exception as e:
				print(f"Error while processing signature: {e}")
				return None  # Handle the error gracefully and return early
		else:
			print("Invalid base64 signature string.")
			return None

	# Process replacements in the document
	for paragraph in doc.paragraphs:
		for placeholder, replacement in replacements.items():
			if placeholder in paragraph.text:
				paragraph.text = paragraph.text.replace(placeholder, replacement)

	# Return the edited document as a BytesIO object
	output_io = BytesIO()
	doc.save(output_io)
	output_io.seek(0)  # Rewind the BytesIO object to the beginning
	return output_io

def send_email(subject, body, to_email, attachments):
	"""Send an email with attachments using Flask-Mail."""
	with app.app_context():  # Ensure the app context is available
		msg = Message(subject, recipients=[to_email], body=body)

		# Attach each file in the attachments list (which are now in-memory)
		for attachment_name, attachment_io in attachments:
			if attachment_io is not None:  # Check if the attachment is valid
				msg.attach(
					attachment_name, 
					'application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
					attachment_io.read()
				)
			else:
				print(f"Warning: Attachment {attachment_name} is None and will not be included.")

		# Send the email
		mail.send(msg)

def submit_form(C_UniqueID):
	"""Handle form submission and send edited documents via email without saving to disk."""
	print("function submit form called")

	activity_mapping = {
		"ZIPLINE": "ZIPLINE",
		"WA": "WATER ABSEILING",
		"LOF": "LEAP OF FAITH",
		"WCLIMG": "WALL CLIMBING"
	}

	address_parts = [
		session.get('lineaddress', ''),
		session.get('postcode', ''),
		session.get('state', ''),
		session.get('country', '')
	]

	activities_info = session.get('activities', [])
	activity_parts = [
		activity_mapping[activity['AN']]  # Map short code to full name
		for activity in activities_info if activity['AN'] in activity_mapping
	]

	activity_parts = [
		f"{activity[0]}"  # Using formatted string for each activity name
		for activity in activities_info
	]
	print("Address components retrieved")

	# Filter out any empty strings
	address = ', '.join(part for part in address_parts if part).strip()

	# Define your replacements for the first document
	form_replacements = {
		"<<NAME>>": session.get("fullname"),
		"<<NRIC>>": session.get("NRIC"),
		"<<AGE>>": session.get("age"),
		"<<GENDER>>": session.get("gender"),
		"<<EMAIL>>": session.get("email"),
		"<<CONTACT>>": session.get("contact"),
		"<<ADDRESS>>": address,
		"<<NATIONALITY>>": session.get("nationality"),
		"<<RACE>>": session.get("race"),
		"<<UNIQID>>": C_UniqueID
	}
	print("Form replacements prepared")

	# Paths to your original docx files
	form_docx_path = "reform/GGP_Form.docx"  # Path to the first document
	e_cert_docx_path = "e-certs/GGPCertificate.docx"  # Path to the second document
	print("Paths set up for document editing")

	# Edit the documents in memory
	edited_form_docx_io = edit_docx_in_memory(form_docx_path, form_replacements)
	print("Form document edited")

	formatted_activity = '\n'.join(activity_parts)

	# Define your replacements for the second document
	e_cert_replacements = {
		"<<NAME>>": session.get("fullname"),
		"<<NRIC>>": session.get("NRIC"),
		"<<ACTIVITY>>": formatted_activity
	}
	print("E Cert replacements prepared")

	edited_e_cert_docx_io = edit_docx_in_memory(e_cert_docx_path, e_cert_replacements)
	print("E cert document edited")

	# Check if both documents were edited successfully
	if edited_form_docx_io is None or edited_e_cert_docx_io is None:
		print("Error: One or more documents failed to generate.")
		return jsonify({"error": "Failed to generate documents."}), 500

	fullname = session.get("fullname")
	subject = "Confirmation of Your Indemnity Form Submission"

	# Check if health_fields is empty or consists of only whitespace
	if not session.get("health_fields", "").strip():
		body = (
			f"Dear {fullname},\n\n"
			"Thank you for submitting your online indemnity form for the upcoming adventure activity. "
			"We want to assure you that your personal information is securely stored in line with our data protection policy.\n\n"
			"Our goal is to provide you with the best possible experience, and we are excited to have you join us for this adventure. "
			"Should you have any questions or need further assistance, please don't hesitate to reach out.\n\n"
			"We look forward to seeing you soon!\n\n"
			"Best regards,\n"
			"Glamping Park Travel And Tour Sdn Bhd"
		)
	else:
		body = (
			f"Dear {fullname},\n\n"
			"Thank you for submitting your online indemnity form.\n"
			"We acknowledge that you have declared medical conditions that may require special attention during the upcoming adventure activity.\n"
			"Please be assured that this information is safe under our Data Protection Policy and has been shared with the trip leader to ensure appropriate care and precaution.\n\n"
			"Please see our trip leader for a short assessment on your condition. "
			"Failure to do so will hinder our effort to give you the best experience you deserve and for safety purposes.\n"
			"If there are any updates or further information you need to provide, please seek attention from our Trip Leader.\n\n"
			"We look forward to seeing you on the adventure!\n\n"
			"Best regards,\n"
			"Glamping Park Travel And Tour Sdn Bhd"
		)
	send_email(
		subject,
		body,
		session.get("email"),  # Use the email from the form data
		[("GGP_Form.docx", edited_form_docx_io), ("GGPCertificate.docx", edited_e_cert_docx_io)]  # Attach both in-memory docs
	)
	
	print("Email sent successfully")

	return jsonify({"message": "Documents sent successfully!"}), 200


if __name__ == "__main__":
	app.run(debug=True)